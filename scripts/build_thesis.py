"""Build the renewed thesis as a styled, data-backed DOCX.

The original 2025 PDF remains untouched. This builder reads only versioned analysis
outputs and publication figures, making every reported number reproducible.
"""

from __future__ import annotations

import argparse
import json
from collections.abc import Iterable, Sequence
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs" / "latest"
FIGURES = OUTPUTS / "figures"
TABLES = OUTPUTS / "tables"
DEFAULT_OUTPUT = ROOT / "docs" / "Atakum_Konut_Fiyatlari_Yenilenmis_Tez.docx"

NAVY = "17324D"
BLUE = "2F80ED"
TEAL = "128C86"
GOLD = "B07A18"
INK = "172B3A"
MUTED = "667788"
LIGHT_BLUE = "EAF2FA"
LIGHT_GRAY = "F3F6F8"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9070


def load_json(filename: str) -> dict:
    return json.loads((OUTPUTS / filename).read_text(encoding="utf-8"))


def set_run_font(
    run,
    *,
    name: str = "Times New Roman",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Sütun genişlikleri {TABLE_WIDTH_DXA} DXA toplamına eşit olmalıdır")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "110")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)
    set_run_font(run, name="Arial", size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (15, NAVY, 18, 10),
        "Heading 2": (12.5, BLUE, 14, 7),
        "Heading 3": (11.5, TEAL, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = True

    if "Kaynakça" not in styles:
        reference = styles.add_style("Kaynakça", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = styles["Kaynakça"]
    reference.font.name = "Times New Roman"
    reference._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    reference._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    reference.font.size = Pt(10)
    reference.paragraph_format.left_indent = Cm(0.75)
    reference.paragraph_format.first_line_indent = Cm(-0.75)
    reference.paragraph_format.line_spacing = 1.15
    reference.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("ATAKUM KONUT İLANLARI • YENİLENMİŞ ANALİZ")
    set_run_font(run, name="Arial", size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    props = doc.core_properties
    props.title = (
        "Samsun Atakum İlçesindeki Konut Fiyatlarına İstatistiksel Yaklaşım ve Fiyat Tahmini"
    )
    props.subject = "Bitirme projesi - yenilenmiş analiz sürümü"
    props.author = "Yunus Emre Büyükgüler; Kadir Ertürk"
    props.keywords = "Atakum, konut ilanı, hedonik model, makine öğrenmesi, fiyat tahmini"


def add_para(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    italic: bool = False,
    keep_with_next: bool = False,
    first_line: bool = True,
) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.first_line_indent = Cm(0.75) if first_line else Cm(0)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        body = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(body, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)


def add_heading(doc: Document, text: str, level: int, *, new_page: bool = False) -> None:
    paragraph = doc.add_heading(text, level=level)
    if new_page:
        paragraph.paragraph_format.page_break_before = True


def add_equation(doc: Document, equation: str, number: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"{equation}    ({number})")
    set_run_font(run, name="Cambria Math", size=11, italic=True, color=NAVY)


def add_table(
    doc: Document,
    caption: str,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths: Sequence[int],
    *,
    source: str = "Kaynak: Yazarların hesaplamaları.",
) -> None:
    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_paragraph.paragraph_format.keep_with_next = True
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, name="Arial", size=9.5, bold=True, color=NAVY)

    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, name="Arial", size=8.5, bold=True, color=WHITE)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    set_run_font(run, name="Arial", size=8.2, color=INK)
    set_table_geometry(table, widths)

    source_paragraph = doc.add_paragraph()
    source_paragraph.paragraph_format.space_before = Pt(4)
    source_paragraph.paragraph_format.space_after = Pt(8)
    run = source_paragraph.add_run(source)
    set_run_font(run, name="Arial", size=8, italic=True, color=MUTED)


def add_figure(
    doc: Document,
    filename: str,
    caption: str,
    *,
    width_cm: float = 15.1,
    source: str = "Kaynak: Yazarların hesaplamaları; veriseti.xlsx.",
) -> None:
    image_path = FIGURES / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    picture = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(". ", 1)[-1])

    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.paragraph_format.keep_with_next = True
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, name="Arial", size=9, bold=True, color=NAVY)
    source_paragraph = doc.add_paragraph()
    source_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_paragraph.paragraph_format.space_after = Pt(9)
    source_run = source_paragraph.add_run(source)
    set_run_font(source_run, name="Arial", size=7.8, italic=True, color=MUTED)


def tr_num(value: float, digits: int = 0) -> str:
    formatted = f"{value:,.{digits}f}"
    return formatted.replace(",", "_").replace(".", ",").replace("_", ".")


def tl(value: float) -> str:
    return f"{tr_num(value)} TL"


def pct(value: float, digits: int = 1) -> str:
    return f"%{tr_num(value * 100, digits)}"


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution.paragraph_format.space_after = Pt(28)
    for line in (
        "T.C.",
        "ONDOKUZ MAYIS ÜNİVERSİTESİ",
        "FEN FAKÜLTESİ",
        "İSTATİSTİK BÖLÜMÜ",
    ):
        run = institution.add_run(line)
        set_run_font(run, name="Arial", size=12, bold=True, color=NAVY)
        run.add_break()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(10)
    set_run_font(
        kicker.add_run("BİTİRME PROJESİ • YENİLENMİŞ ANALİZ SÜRÜMÜ"),
        name="Arial",
        size=9,
        bold=True,
        color=GOLD,
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(
        "Samsun Atakum İlçesindeki Konut Fiyatlarına\nİstatistiksel Yaklaşım ve Fiyat Tahmini"
    )
    set_run_font(run, name="Arial", size=19, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(48)
    set_run_font(
        subtitle.add_run("Tekrarlanabilir istatistiksel analiz ve makine öğrenmesi uygulaması"),
        name="Arial",
        size=11,
        italic=True,
        color=MUTED,
    )

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(28)
    run = authors.add_run("Yunus Emre BÜYÜKGÜLER\nKadir ERTÜRK")
    set_run_font(run, name="Arial", size=12, bold=True, color=INK)

    advisor = doc.add_paragraph()
    advisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    advisor.paragraph_format.space_after = Pt(56)
    set_run_font(
        advisor.add_run("Danışman\nÖğretim Görevlisi Umut YAMAK"), name="Arial", size=11, color=INK
    )

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("Samsun • 2026"), name="Arial", size=10, bold=True, color=MUTED)
    date.add_run().add_break(WD_BREAK.PAGE)


def add_front_matter(doc: Document, summary: dict) -> None:
    price = summary["price"]
    model = summary["best_model"]
    audit = summary["data_audit"]

    add_heading(doc, "SÜRÜM NOTU", 1)
    add_para(
        doc,
        "Bu belge, 2025 tarihli bitirme projesinin konu, yazar, danışman ve temel bölüm yapısını "
        "koruyan; veri hattını, istatistiksel çıkarımları ve tahmin değerlendirmesini yeniden "
        "üreten sürümdür. Orijinal bitirmetezi.pdf dosyası tarihsel kayıt olarak değiştirilmeden "
        "saklanmıştır. Yenilenmiş sürümdeki sayısal sonuçların tamamı repodaki veriseti.xlsx "
        "dosyası ve sürümlenmiş Python kodundan üretilebilir.",
    )
    add_para(
        doc,
        "Ham veri dosyası repoda bulunmadığından, ilk 2.836 kayıttan çıkarıldığı bildirilen 555 "
        "gözlemin nedenleri bu çalışmada geriye dönük olarak yeniden sınıflandırılmamıştır. "
        "Repodaki 2.281 satır önceden temizlenmiş başlangıç verisi kabul edilmiş; bundan sonraki "
        "dokuz yapısal ret ve ilan-kimliği temelli tekilleştirme ayrıca raporlanmıştır.",
    )
    doc.add_page_break()

    add_heading(doc, "ÖZET", 1)
    abstract = (
        "Bu çalışmada Samsun ili Atakum ilçesinde çevrim içi emlak ilanlarından derlenen konut "
        "özellikleri ile ilan fiyatları arasındaki ilişkiler incelenmiş ve fiyat tahmini için "
        "tekrarlanabilir bir makine öğrenmesi hattı kurulmuştur. Bildirilen ham örneklem 2.836 "
        f"kayıttır. Repoda bulunan {tr_num(audit['repository_rows'])} satır üzerinde yapılan "
        f"yapısal denetimde {tr_num(audit['malformed_rows'])} bozuk kayıt ayrılmış; aynı ilanların "
        f"farklı tarihlerdeki görüntüleri arasından son gözlem tutulduğunda {tr_num(audit['latest_snapshot_rows'])} "
        "tekil ilan elde edilmiştir. Ana analiz birimi bu tekil son görüntülerdir. İlan fiyatının "
        f"ortalaması {tl(price['mean'])}, medyanı {tl(price['median'])} ve medyan brüt metrekare "
        f"fiyatı {tl(price['median_price_per_gross_m2'])} olarak hesaplanmıştır. En az 20 gözlem "
        "içeren mahallelerin fiyat dağılımları Kruskal-Wallis testiyle karşılaştırılmış ve güçlü "
        f"bir farklılık bulunmuştur (H={tr_num(summary['neighborhood_test']['h_statistic'], 2)}, "
        f"p<0,001; ε²={tr_num(summary['neighborhood_test']['epsilon_squared'], 3)}). Yarı logaritmik "
        "hedonik model HC3 dayanıklı standart hatalarla tahmin edilmiş; düzeltilmiş R² "
        f"{tr_num(summary['hedonic_model']['adjusted_r_squared'], 3)} olmuştur. Beş aday tahmin "
        f"yaklaşımı arasında eğitim verisindeki beş katlı çapraz doğrulama MAE'sine göre {model['name']} "
        f"seçilmiştir. Kilitli test kümesinde MAE {tl(model['locked_test']['test_mae'])}, R² "
        f"{tr_num(model['locked_test']['test_r2'], 3)} ve MAPE {pct(model['locked_test']['test_mape'])} "
        "bulunmuştur. Satır bazlı rastgele bölmenin, aynı ilana ait tekrar görüntüler nedeniyle "
        "daha iyimser sonuç verdiği gösterilmiştir. Bulgular gerçekleşmiş satış değerlerini değil "
        "ilan fiyatlarını açıklamakta; model resmi ekspertiz aracı olarak önerilmemektedir."
    )
    add_para(doc, abstract)
    add_para(
        doc,
        "Anahtar Sözcükler: konut ilan fiyatı, Atakum, hedonik fiyat modeli, Extra Trees, veri "
        "sızıntısı, çapraz doğrulama.",
        bold_lead="Anahtar Sözcükler:",
        first_line=False,
    )
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1)
    add_para(
        doc,
        "This study examines the associations between online residential listing attributes and "
        "asking prices in Atakum, Samsun, and establishes a reproducible machine-learning pipeline "
        "for price prediction. The reported raw sample contains 2,836 records. Structural checks "
        f"on the {audit['repository_rows']:,} repository rows rejected {audit['malformed_rows']} malformed "
        f"records, while retaining the latest snapshot per listing produced {audit['latest_snapshot_rows']:,} "
        "unique listings for the primary analyses. Median asking price was TRY "
        f"{price['median']:,.0f}, and median asking price per gross square metre was TRY "
        f"{price['median_price_per_gross_m2']:,.0f}. Price distributions differed across eligible "
        "neighbourhoods in a Kruskal-Wallis test. A semi-log hedonic specification used HC3 robust "
        "standard errors. Among five candidates, Extra Trees was selected by five-fold cross-validated "
        f"MAE on the training set and achieved a locked-test MAE of TRY {model['locked_test']['test_mae']:,.0f}, "
        f"R²={model['locked_test']['test_r2']:.3f}, and MAPE={model['locked_test']['test_mape'] * 100:.1f}%. "
        "Sensitivity checks show that randomly splitting repeated listing snapshots yields optimistic "
        "performance. Results concern asking prices rather than completed transactions and should not "
        "be interpreted as a formal appraisal.",
    )
    add_para(
        doc,
        "Keywords: housing asking price, Atakum, hedonic price model, Extra Trees, data leakage, "
        "cross-validation.",
        bold_lead="Keywords:",
        first_line=False,
    )
    doc.add_page_break()

    add_heading(doc, "ÖNSÖZ VE TEŞEKKÜR", 1)
    add_para(
        doc,
        "Bu çalışma, Ondokuz Mayıs Üniversitesi Fen Fakültesi İstatistik Bölümünde yürütülen "
        "bitirme projesinin yenilenmiş analiz sürümüdür. Konunun belirlenmesinden çalışmanın "
        "tamamlanmasına kadar sağladığı akademik yönlendirme için danışmanımız Öğretim Görevlisi "
        "Umut YAMAK'a teşekkür ederiz.",
    )
    add_para(
        doc,
        "Yenileme sürecinde temel amaç, önceki çalışmanın kapsamını genişletmekten çok mevcut "
        "veriyi daha izlenebilir, tekrar üretilebilir ve yöntemsel olarak daha temkinli bir biçimde "
        "sunmaktır. Bu nedenle elde bulunmayan ham kayıtlar hakkında yeni varsayımlar kurulmamış, "
        "bütün ek kararlar kod ve çıktı dosyalarıyla belgelenmiştir.",
    )
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.paragraph_format.space_before = Pt(22)
    set_run_font(
        signature.add_run("Yunus Emre BÜYÜKGÜLER • Kadir ERTÜRK"),
        name="Arial",
        size=10,
        bold=True,
        color=NAVY,
    )
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "İÇİNDEKİLER", 1)
    entries = [
        ("1. GİRİŞ", "9"),
        ("2. KAVRAMSAL ÇERÇEVE VE LİTERATÜR", "12"),
        ("3. VERİ VE YÖNTEM", "13"),
        ("4. BULGULAR", "17"),
        ("5. TARTIŞMA", "26"),
        ("6. SONUÇ VE ÖNERİLER", "28"),
        ("KAYNAKÇA", "29"),
        ("EKLER", "30"),
    ]
    for label, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        label_run = paragraph.add_run(label)
        set_run_font(label_run, name="Arial", size=10.5, bold=True, color=NAVY)
        paragraph.add_run("\t")
        page_run = paragraph.add_run(page)
        set_run_font(page_run, name="Arial", size=10.5, bold=True, color=MUTED)
    add_para(
        doc,
        "Not: Sayfa numaraları yenilenmiş PDF sürümünün bölüm başlangıçlarını gösterir.",
        italic=True,
        first_line=False,
    )

    add_heading(doc, "ŞEKİLLER DİZİNİ", 2)
    figures = [
        "Şekil 3.1. Örneklem ve veri hazırlama akışı",
        "Şekil 4.1. Tekil ilan fiyatlarının dağılımı",
        "Şekil 4.2. Mahalle bazında medyan ilan fiyatları",
        "Şekil 4.3. Oda tipine göre fiyat profili",
        "Şekil 4.4. Sayısal değişkenler için Spearman korelasyonları",
        "Şekil 4.5. Aday modellerin çapraz doğrulama performansı",
        "Şekil 4.6. Kilitli testte gerçekleşen ve tahmin edilen fiyatlar",
        "Şekil 4.7. Kilitli test artıklarının dağılımı",
        "Şekil 4.8. Permütasyon özelliği önemleri",
        "Şekil 4.9. Değerlendirme protokolü duyarlılığı",
    ]
    for figure in figures:
        add_para(doc, figure, first_line=False)

    add_heading(doc, "TABLOLAR DİZİNİ", 2)
    tables = [
        "Tablo 3.1. Veri soy ağacı ve örneklem kararları",
        "Tablo 3.2. Analizde kullanılan değişken grupları",
        "Tablo 3.3. Model değerlendirme protokolleri",
        "Tablo 4.1. Temel tanımlayıcı istatistikler",
        "Tablo 4.2. Medyan fiyatı en yüksek mahalleler",
        "Tablo 4.3. İkili grup karşılaştırmaları",
        "Tablo 4.4. Hedonik modelin seçili katsayıları",
        "Tablo 4.5. Tahmin modeli karşılaştırması",
        "Tablo 4.6. Değerlendirme protokolü duyarlılığı",
    ]
    for table_name in tables:
        add_para(doc, table_name, first_line=False)
    doc.add_page_break()

    add_heading(doc, "SİMGELER VE KISALTMALAR", 1)
    add_table(
        doc,
        "Kullanılan simge ve kısaltmalar",
        ["Gösterim", "Açıklama"],
        [
            ("ε²", "Kruskal-Wallis etki büyüklüğü (epsilon kare)"),
            ("HC3", "Değişen varyansa dayanıklı kovaryans tahmincisi"),
            ("MAE", "Ortalama mutlak hata"),
            ("MAPE", "Ortalama mutlak yüzde hata"),
            ("RMSE", "Kök ortalama kare hata"),
            ("R²", "Belirleme katsayısı"),
            ("CV", "Çapraz doğrulama"),
            ("IQR", "Çeyrekler arası açıklık"),
        ],
        [1700, 7370],
        source="",
    )


def add_introduction(doc: Document) -> None:
    add_heading(doc, "1. GİRİŞ", 1, new_page=True)
    add_heading(doc, "1.1. Problem ve araştırma bağlamı", 2)
    add_para(
        doc,
        "Konut, barınma işlevinin yanında yüksek parasal değeri, uzun kullanım süresi ve konuma "
        "bağımlılığı nedeniyle heterojen bir ekonomik maldır. Aynı mahallede bulunan iki konut; "
        "alan, bina yaşı, kat, ısıtma sistemi, site durumu ve başka nitelikler bakımından farklılaşır. "
        "Bu heterojenlik, fiyat karşılaştırmasının yalnızca ortalama ya da tek bir metrekare fiyatına "
        "indirgenmesini güçleştirir. Hedonik yaklaşım, toplam fiyatı gözlenebilen özelliklerin birlikte "
        "oluşturduğu bir sonuç olarak ele alır (Rosen, 1974).",
    )
    add_para(
        doc,
        "Çevrim içi ilan platformları yerel piyasayı ayrıntılı biçimde gözlemlemek için zengin veri "
        "üretse de bu veriler satış sözleşmelerini değil satıcıların talep ettiği fiyatları gösterir. "
        "Aynı ilan farklı tarihlerde tekrar görüntülenebilir, bazı alanlar eksik bırakılabilir ve ilan "
        "fiyatı pazarlıkla oluşan nihai satış bedelinden ayrışabilir. Dolayısıyla ilan verisiyle yapılan "
        "bir çalışma hem gözlem birimini doğru tanımlamalı hem de tahmin performansını tekrar kayıtların "
        "yaratabileceği bilgi sızıntısına karşı sınamalıdır.",
    )

    add_heading(doc, "1.2. Araştırmanın amacı", 2)
    add_para(
        doc,
        "Çalışmanın amacı, Atakum ilçesinde derlenen satılık konut ilanlarını şeffaf bir veri hattıyla "
        "incelemek, fiyat dağılımını ve mahalleler arası farklılıkları betimlemek, hedonik bir modelle "
        "koşullu ilişkileri değerlendirmek ve farklı makine öğrenmesi yöntemlerinin genellenebilir "
        "tahmin başarısını karşılaştırmaktır. Yenilenmiş sürüm, yalnızca daha yüksek bir R² elde etmeyi "
        "değil, bu değerin hangi örneklem ve test protokolü altında üretildiğini göstermeyi hedefler.",
    )

    add_heading(doc, "1.3. Araştırma soruları ve hipotezler", 2)
    questions = [
        ("AS1", "Tekil son ilanlarda fiyat ve fiziksel özelliklerin dağılımı nasıldır?"),
        ("AS2", "Yeterli gözlem içeren mahallelerin ilan fiyatı dağılımları farklılaşmakta mıdır?"),
        (
            "AS3",
            "Alan, bina yaşı, kat ve konum değişkenleri kontrol edildiğinde hangi özellikler ilan fiyatıyla ilişkilidir?",
        ),
        (
            "AS4",
            "Aday tahmin modellerinden hangisi eğitim verisindeki çapraz doğrulamada en düşük hatayı verir?",
        ),
        (
            "AS5",
            "Tekrarlanan ilan görüntülerinin satır bazlı bölünmesi test performansını ne ölçüde iyimserleştirir?",
        ),
    ]
    add_table(
        doc,
        "Araştırma soruları",
        ["Kod", "Soru"],
        questions,
        [1100, 7970],
        source="Kaynak: Araştırma tasarımı.",
    )
    add_para(
        doc,
        "Mahalle karşılaştırmasının yokluk hipotezi, uygun mahallelerin fiyat dağılımlarının aynı "
        "olduğudur. İkili nitelikler için de iki grubun fiyat dağılımlarının eşit olduğu hipotezi "
        "kurulmuştur. Çoklu ikili testlerde aile düzeyi hata olasılığını sınırlamak için Holm "
        "düzeltmesi uygulanmıştır (Holm, 1979).",
    )

    add_heading(doc, "1.4. Kapsam ve katkı", 2)
    add_para(
        doc,
        "Coğrafi kapsam Samsun ili Atakum ilçesi, ürün kapsamı satılık konut ilanlarıdır. Repo "
        "verisindeki gözlem tarihleri 23 Haziran 2024 ile 19 Mayıs 2025 arasındadır; yoğunluk Nisan-Mayıs "
        "2025 dönemindedir. Çalışma zamansal bir fiyat endeksi üretmemekte ve ilçe dışına istatistiksel "
        "genelleme yapmamaktadır. Katkı, yerel bir veri seti üzerinde veri soy ağacı, dayanıklı çıkarım, "
        "kilitli test kümesi ve alternatif bölme protokollerini tek bir tekrarlanabilir iş akışında "
        "birleştirmesidir.",
    )


def add_literature(doc: Document) -> None:
    add_heading(doc, "2. KAVRAMSAL ÇERÇEVE VE LİTERATÜR", 1, new_page=True)
    add_heading(doc, "2.1. Hedonik fiyat yaklaşımı", 2)
    add_para(
        doc,
        "Hedonik fiyat teorisinde farklılaştırılmış bir malın piyasa fiyatı, o malın niteliklerinin "
        "oluşturduğu bileşimin bir fonksiyonu olarak düşünülür. Rosen (1974), gözlenen ürün fiyatları ve "
        "özellik miktarları üzerinden örtük özellik fiyatlarının analiz edilebileceği kuramsal çerçeveyi "
        "geliştirmiştir. Konut araştırmalarında büyüklük, konum, yaş ve donatı değişkenlerinin aynı modelde "
        "yer alması bu nedenle yalnızca tahmin amacı taşımaz; fiyat farklılıklarının hangi gözlenen "
        "bileşimlerle birlikte ortaya çıktığını açıklamaya da yardımcı olur.",
    )
    add_para(
        doc,
        "Yarı logaritmik biçimde bağımlı değişkenin logaritması modellenir. Sürekli bir açıklayıcıdaki "
        "bir birimlik değişimin fiyatla yüzdesel birlikteliği yaklaşık olarak katsayı üzerinden, kukla "
        "değişkenlerde ise exp(β)-1 dönüşümüyle yorumlanabilir. Bununla birlikte gözlemsel ilan verisinde "
        "katsayılar nedensel etki değildir. Ölçülmeyen manzara, konutun iç kalite düzeyi, cephe ve mikro "
        "konum gibi özellikler hem fiyatla hem de modele giren değişkenlerle ilişkili olabilir.",
    )

    add_heading(doc, "2.2. Türkiye'de konut fiyatı araştırmaları", 2)
    add_para(
        doc,
        "Selim (2009), Türkiye 2004 Hanehalkı Bütçe Anketi verisiyle hedonik regresyon ve yapay sinir "
        "ağını karşılaştırmış; doğrusal olmayan yapıyı yakalayan sinir ağının tahmin performansında üstün "
        "olabileceğini göstermiştir. Bu bulgu, konut nitelikleri ile fiyat arasındaki ilişkinin yalnızca "
        "doğrusal bir formda ele alınmaması gerektiğine işaret eder. Ancak ulusal anket verisiyle yerel "
        "ilan verisinin hedef değişkeni, örnekleme mekanizması ve zaman kapsamı farklıdır; bu nedenle "
        "sonuçların doğrudan performans karşılaştırması olarak okunması uygun değildir.",
    )

    add_heading(doc, "2.3. Ağaç toplulukları ve düzenlileştirme", 2)
    add_para(
        doc,
        "Rastgele Orman, farklı bootstrap örnekleri ve rastgele özellik alt kümeleri üzerinde kurulan "
        "ağaçların tahminlerini birleştirerek tek bir ağacın yüksek varyansını azaltmayı amaçlar "
        "(Breiman, 2001). Extra Trees yöntemi ise özellik ve kesme noktası seçimindeki rassallığı daha "
        "da artırır (Geurts, Ernst ve Wehenkel, 2006). Gradient boosting yaklaşımı, ardışık zayıf "
        "öğrenicilerin kayıp fonksiyonu yönünde eklenmesiyle esnek fonksiyon tahmini üretir (Friedman, "
        "2001). Ridge regresyon, katsayıların karesine ceza ekleyerek çoklu bağlantı ve yüksek varyans "
        "karşısında daha istikrarlı bir doğrusal kestirim sağlar (Hoerl ve Kennard, 1970).",
    )
    add_para(
        doc,
        "Bu çalışmada tek bir karmaşık modele öncelik vermek yerine, medyan temel modelden doğrusal "
        "düzenlileştirmeye ve üç farklı ağaç topluluğuna uzanan adaylar aynı eğitim/test çerçevesinde "
        "karşılaştırılmıştır. Uygulama scikit-learn ekosistemiyle hazırlanmıştır (Pedregosa ve diğerleri, "
        "2011). Kategorik kodlama, eksik değer tamamlama ve ölçekleme işlemleri model boru hattının içinde "
        "tutulmuş; böylece test gözlemlerinin eğitim ön işlemlerine bilgi taşıması engellenmiştir.",
    )

    add_heading(doc, "2.4. Veri sızıntısı ve değerlendirme tasarımı", 2)
    add_para(
        doc,
        "Tahmin modelinde bilgi sızıntısı, gerçek kullanım anında bulunmayacak bilginin eğitim sürecine "
        "girmesi veya aynı birime ait çok benzer kayıtların eğitim ve test tarafında birlikte yer alması "
        "durumunda ortaya çıkabilir. İlan verisinde aynı ilan numarasıyla birden çok tarih görüntüsü "
        "bulunması, satır bazlı rastgele bölmede modelin test ilanına çok yakın bir eğitim örneği görmesine "
        "neden olur. Bu durum genellenebilirliği değil kısmi hatırlamayı ölçebilir.",
    )
    add_para(
        doc,
        "Yenilenmiş tasarım bu sorunu dört ayrı görünümle ele alır: tekil son ilanlarda kilitli rastgele "
        "test, ileri tarihlerdeki ilanları ayıran kronolojik kontrol, tüm görüntülerde iyimserlik riski "
        "taşıyan satır bazlı bölme ve ilan numarasını iki tarafa ayırmayan grup bölmesi. Bu karşılaştırma, "
        "tek bir başarı oranından daha açıklayıcı bir performans aralığı sunar.",
    )

    add_heading(doc, "2.5. Araştırma boşluğu", 2)
    add_para(
        doc,
        "Yerel konut çalışmaları çoğu zaman çok sayıda algoritma ve grafik sunsa da ham veriden raporlanan "
        "sonuca giden kararların yeniden üretimi sınırlı kalabilir. Bu çalışma; silinen kayıt sayısını, "
        "yapısal retleri, tekrar ilanları, model seçimini ve sağlamlık kontrollerini ayrı katmanlarda "
        "belgeleyerek bu boşluğa uygulamalı bir yanıt vermektedir. Sonuçların tamamı CSV/JSON tabloları, "
        "sabit rastgelelik tohumu ve çalıştırılabilir kodla birlikte sunulmuştur.",
    )


def add_methodology(doc: Document, summary: dict) -> None:
    audit = summary["data_audit"]
    add_heading(doc, "3. VERİ VE YÖNTEM", 1, new_page=True)
    add_heading(doc, "3.1. Araştırma tasarımı ve gözlem birimi", 2)
    add_para(
        doc,
        "Araştırma nicel, gözlemsel ve kesitsel ağırlıklı bir uygulamadır. Veriler Sahibinden.com "
        "üzerindeki ilanlardan manuel olarak derlenmiştir. Bir satır, ilan sayfasının belirli bir "
        "tarihteki görüntüsünü temsil eder; bir konutun gerçekleşmiş satış işlemi değildir. Ana analiz "
        "birimi ilan numarasıyla tanımlanan tekil ilanın en son gözlemidir. Bu karar, aynı ilanı birden "
        "fazla kez ağırlıklandırmayı ve tahmin testinde benzer görüntülerin iki tarafa dağılmasını azaltır.",
    )

    add_heading(doc, "3.2. Veri soy ağacı ve kalite kontrolleri", 2)
    lineage_rows = [
        (
            "Bildirilen ham örneklem",
            tr_num(audit["reported_raw_rows"]),
            "Ham dosya repoda bulunmuyor",
        ),
        (
            "Önceden temizlenmiş repo verisi",
            tr_num(audit["repository_rows"]),
            f"{tr_num(audit['previously_removed_rows'])} kayıt önceki süreçte çıkarılmış",
        ),
        (
            "Yapısal olarak geçerli görüntüler",
            tr_num(audit["structurally_valid_rows"]),
            f"{tr_num(audit['malformed_rows'])} sütun kaymalı/geçersiz kayıt ayrıldı",
        ),
        (
            "Tekil ilanların son görüntüsü",
            tr_num(audit["latest_snapshot_rows"]),
            "İlan numarası başına en güncel kayıt",
        ),
    ]
    add_table(
        doc,
        "Tablo 3.1. Veri soy ağacı ve örneklem kararları",
        ["Aşama", "n", "Açıklama"],
        lineage_rows,
        [3300, 1000, 4770],
    )
    add_figure(
        doc, "01_veri_akisi.png", "Şekil 3.1. Örneklem ve veri hazırlama akışı", width_cm=14.6
    )
    add_para(
        doc,
        f"Repodaki {tr_num(audit['repository_rows'])} satırın {tr_num(audit['exact_duplicate_rows'])}'i "
        "birebir tekrardır. İlan numarası tekrarlanan grupların içindeki toplam görüntü sayısı "
        f"{tr_num(audit['repeated_snapshot_rows'])}, fiyatı en az bir kez değişen ilan sayısı "
        f"{tr_num(audit['listings_with_price_change'])}'dir. Yapısal olarak geçerli "
        f"{tr_num(audit['structurally_valid_rows'])} görüntüden tekil son ilana geçişte "
        f"{tr_num(audit['structurally_valid_rows'] - audit['latest_snapshot_rows'])} eski görüntü ana "
        "örneklem dışında bırakılmıştır. Bu kayıtlar hatalı oldukları için değil, gözlem birimini ilan "
        "düzeyinde sabitlemek için ayrılmıştır.",
    )

    add_heading(doc, "3.3. Değişkenler", 2)
    variable_rows = [
        ("Hedef", "Fiyat (TL)", "İlanda talep edilen toplam fiyat"),
        ("Alan", "Brüt m², Net m²", "Konut büyüklüğü"),
        ("Yapı", "Bina yaşı, kat, kat sayısı, banyo, oda", "Fiziksel özellikler"),
        ("Konum", "Mahalle", "İlan adresinden türetilen alt bölge"),
        ("Donatı", "Isıtma, mutfak, balkon, asansör, otopark", "Konut ve bina donatıları"),
        ("Piyasa", "İlan günü, kimden, takas, kredi uygunluğu", "İlan zamanı ve satış koşulları"),
        ("Yardımcı", "Aidat, eşyalı, kullanım, site, tapu", "İlan niteliği ve kullanım durumu"),
    ]
    add_table(
        doc,
        "Tablo 3.2. Analizde kullanılan değişken grupları",
        ["Grup", "Değişkenler", "İşlev"],
        variable_rows,
        [1500, 4300, 3270],
    )

    add_heading(doc, "3.4. Veri hazırlama", 2)
    add_para(
        doc,
        "İlan tarihleri Türkçe ay adlarını dikkate alan bir dönüşümle ISO tarihe çevrilmiştir. İlan "
        "numaraları metin olarak standartlaştırılmış, fiyat ve alan sütunları sayısala dönüştürülmüştür. "
        "Pozitif fiyat, pozitif brüt/net alan, geçerli tarih ve 8-12 haneli ilan numarası temel yapısal "
        "koşullardır. Bu koşulları sağlamayan dokuz satır reddedilen_kayitlar.csv dosyasına gerekçesiyle "
        "yazılmıştır. Ek bir istatistiksel aykırı değer silme işlemi uygulanmamıştır.",
    )
    add_para(
        doc,
        "Aidat alanındaki metin ve para birimi ifadeleri sayısal değere dönüştürülmüş; eksik değerler "
        "ham veri üzerinde küresel bir değerle doldurulmamıştır. Modelleme sırasında sayısal eksikler "
        "yalnızca eğitim katındaki medyanla, kategorik eksikler en sık sınıfla tamamlanmıştır. Nadir "
        "kategoriler OneHotEncoder içinde asgari frekans eşiğiyle bir araya getirilmiş ve bilinmeyen test "
        "kategorileri güvenli biçimde ele alınmıştır.",
    )

    add_heading(doc, "3.5. Tanımlayıcı ve çıkarımsal analiz", 2)
    add_para(
        doc,
        "Fiyat ve sayısal değişkenler için gözlem sayısı, eksik değer, ortalama, standart sapma, "
        "çeyrekler, varyans, IQR ve çarpıklık hesaplanmıştır. Fiyatın sağa çarpık yapısı nedeniyle "
        "merkezi eğilim yorumlarında medyan öne çıkarılmıştır. Sayısal değişkenlerin monoton ilişkileri "
        "Spearman sıra korelasyonuyla özetlenmiştir.",
    )
    add_para(
        doc,
        "İki gruplu nitelikler için iki yönlü Mann-Whitney U testi ve sıra-biserial etki büyüklüğü "
        "kullanılmıştır (Mann ve Whitney, 1947). Altı testin p değerleri Holm yöntemiyle düzeltilmiştir. "
        "Mahalle analizi, en az 20 gözlemi bulunan gruplarla Kruskal-Wallis H testi üzerinden yapılmış "
        "ve etki büyüklüğü epsilon kare ile raporlanmıştır (Kruskal ve Wallis, 1952). Bu testlerin "
        "dağılım farkını değerlendirdiği, tek başına konumsal nedensellik göstermediği dikkate alınmıştır.",
    )

    add_heading(doc, "3.6. Hedonik model", 2)
    add_para(
        doc,
        "Hedonik modelde bağımlı değişken doğal logaritmik ilan fiyatıdır. Brüt alan da logaritmik "
        "biçimde modele alınmış; bina yaşı, banyo sayısı, oda sayısı, bulunduğu kat, toplam kat sayısı, "
        "ilan günü ile mahalle, site, asansör, otopark, ısıtma ve satıcı türü göstergeleri kullanılmıştır.",
    )
    add_equation(doc, "ln(Fiyatᵢ) = β₀ + β₁ ln(Brüt m²ᵢ) + β′Xᵢ + εᵢ", "3.1")
    add_para(
        doc,
        "Katsayı belirsizliği değişen varyansa karşı HC3 dayanıklı standart hatalarla hesaplanmıştır. "
        "HC3, yüksek kaldıraçlı gözlemlerin artıklarını 1-hᵢᵢ terimiyle ölçekleyerek sonlu örneklemde "
        "daha temkinli kovaryans tahmini amaçlar (MacKinnon ve White, 1985). Model açıklayıcıdır; kukla "
        "katsayıları referans kategoriye koşullu birliktelik olarak okunmuştur.",
    )

    add_heading(doc, "3.7. Tahmin modelleri ve model seçimi", 2)
    add_para(
        doc,
        "Medyan temel model, Ridge regresyon, Random Forest, Extra Trees ve Histogram Gradient Boosting "
        "aynı özellikler ve logaritmik hedef dönüşümü altında karşılaştırılmıştır. Tekil ilanların yüzde "
        "80'i eğitim, yüzde 20'si kilitli test için ayrılmış; fiyat onluklarına göre yaklaşık tabakalama "
        "yapılmıştır. Model seçimi kilitli test sonuçlarına bakılarak değil, yalnızca eğitim parçasındaki "
        "beş katlı çapraz doğrulama MAE'siyle gerçekleştirilmiştir. Rastgelelik tohumu 42'dir.",
    )

    add_heading(doc, "3.8. Değerlendirme ölçütleri", 2)
    add_equation(doc, "MAE = (1/n) Σ |yᵢ - ŷᵢ|", "3.2")
    add_equation(doc, "RMSE = √[(1/n) Σ (yᵢ - ŷᵢ)²]", "3.3")
    add_equation(doc, "R² = 1 - Σ(yᵢ - ŷᵢ)² / Σ(yᵢ - ȳ)²", "3.4")
    add_equation(doc, "MAPE = (1/n) Σ |(yᵢ - ŷᵢ) / yᵢ|", "3.5")
    add_para(
        doc,
        "MAE doğrudan TL cinsinden tipik mutlak hatayı, RMSE büyük hatalara daha fazla ağırlık veren "
        "karekök hatayı, R² test varyansının açıklanan kısmını ve MAPE göreli hatayı gösterir. MAE için "
        "kilitli test hataları üzerinden 2.000 bootstrap tekrarına dayalı yüzde 95 aralık verilmiştir. "
        "Bu aralık modelin ortalama hata düzeyine ilişkindir; tek bir konut için tahmin aralığı değildir.",
    )

    add_heading(doc, "3.9. Sağlamlık protokolleri", 2, new_page=True)
    protocol_rows = [
        ("Kilitli tekil test", "Tekil son ilanlar", "Genel test performansı"),
        ("Kronolojik test", "Erken tarihler eğitim, ileri tarihler test", "Zamansal aktarım"),
        ("Satır bazlı rastgele", "Tüm geçerli görüntüler", "Sızıntı riskinin gösterimi"),
        ("İlan kimliği gruplu", "Aynı ilan tek tarafta", "Tekrar ilanlardan bağımsız genelleme"),
    ]
    add_table(
        doc,
        "Tablo 3.3. Model değerlendirme protokolleri",
        ["Protokol", "Bölme birimi", "Amaç"],
        protocol_rows,
        [2600, 3400, 3070],
    )

    add_heading(doc, "3.10. Tekrarlanabilirlik ve etik sınırlar", 2)
    add_para(
        doc,
        "Analiz, Python 3.10 ve üzeri sürümlerde çalışacak paket yapısında hazırlanmıştır. Veri işlemleri "
        "pandas, sayısal hesaplamalar NumPy/SciPy, modelleme scikit-learn ve görseller Matplotlib/Seaborn "
        "ile yürütülür. Tek komutla tablo, JSON, şekil ve model çıktıları yeniden oluşturulur. Testler "
        "veri dönüşümü, metrik hesapları ve Streamlit dashboard açılışını denetler.",
    )
    add_para(
        doc,
        "İlan numarası teknik tekilleştirme anahtarıdır; raporun anlatımında bireysel ilan kimlikleri "
        "paylaşılmaz. Dashboard küçük alt gruplarda dikkatli yorum uyarısı verir. Fiyat simülatörü resmi "
        "ekspertiz, kredi kararı, satış garantisi veya yatırım tavsiyesi değildir.",
    )


def add_results(doc: Document, summary: dict) -> None:
    price = summary["price"]
    model = summary["best_model"]
    desc = pd.read_csv(TABLES / "tanimlayici_istatistikler.csv")
    neighborhoods = pd.read_csv(TABLES / "mahalle_ozeti.csv")
    binary = pd.read_csv(TABLES / "ikili_grup_testleri.csv")
    hedonic = pd.read_csv(TABLES / "hedonik_katsayilar_hc3.csv")
    comparison = pd.read_csv(TABLES / "model_karsilastirma.csv")
    protocols = pd.read_csv(TABLES / "protokol_duyarliligi.csv")
    importance = pd.read_csv(TABLES / "ozellik_onemi.csv")

    add_heading(doc, "4. BULGULAR", 1, new_page=True)
    add_heading(doc, "4.1. Fiyat ve fiziksel özelliklerin dağılımı", 2)
    selected = desc.loc[
        desc["variable"].isin(
            ["Fiyat (TL)", "Brüt m²", "Net m²", "Bina Yaşı Ortalama", "Brüt m² Başına Fiyat"]
        )
    ]
    rows = []
    labels = {
        "Fiyat (TL)": "İlan fiyatı (TL)",
        "Brüt m²": "Brüt alan (m²)",
        "Net m²": "Net alan (m²)",
        "Bina Yaşı Ortalama": "Bina yaşı",
        "Brüt m² Başına Fiyat": "Brüt m² fiyatı (TL)",
    }
    for _, row in selected.iterrows():
        rows.append(
            (
                labels[row["variable"]],
                tr_num(row["count"]),
                tr_num(row["mean"], 1),
                tr_num(row["50%"], 1),
                tr_num(row["25%"], 1),
                tr_num(row["75%"], 1),
            )
        )
    add_table(
        doc,
        "Tablo 4.1. Temel tanımlayıcı istatistikler",
        ["Değişken", "n", "Ortalama", "Medyan", "Q1", "Q3"],
        rows,
        [2600, 900, 1450, 1450, 1335, 1335],
    )
    add_para(
        doc,
        f"Tekil ilanlarda fiyat {tl(price['minimum'])} ile {tl(price['maximum'])} arasında değişmektedir. "
        f"Ortalama {tl(price['mean'])}, medyan {tl(price['median'])} olduğundan dağılımın yüksek fiyatlı "
        "ilanlar yönünde sağa çarpık olduğu görülür. Fiyat değişkeninin çarpıklık katsayısı 0,947'dir. "
        f"Medyan brüt metrekare ilan fiyatı {tl(price['median_price_per_gross_m2'])}'dir.",
    )
    add_figure(doc, "02_fiyat_dagilimi.png", "Şekil 4.1. Tekil ilan fiyatlarının dağılımı")
    add_para(
        doc,
        "Dağılım grafiği, 2-4 milyon TL bandındaki yoğunluğu ve üst taraftaki uzun kuyruğu birlikte "
        "göstermektedir. Uç fiyatlar otomatik olarak silinmemiştir; çünkü veri setinin tanımlı aralığında "
        "geçerli ilanlar olabilirler. Dayanıklı medyan/IQR özetleri ve logaritmik hedef dönüşümü, bu "
        "asimetriyi analiz içinde yönetmek için kullanılmıştır.",
    )

    add_heading(doc, "4.2. Mahalle düzeyinde farklılıklar", 2)
    top = neighborhoods.head(8)
    neighborhood_rows = [
        (
            row["Mahalle"],
            tr_num(row["listing_count"]),
            tl(row["median_price"]),
            tl(row["median_price_per_gross_m2"]),
            tr_num(row["median_gross_m2"], 1),
        )
        for _, row in top.iterrows()
    ]
    add_table(
        doc,
        "Tablo 4.2. Medyan fiyatı en yüksek mahalleler",
        ["Mahalle", "n", "Medyan fiyat", "Medyan m² fiyatı", "Medyan brüt m²"],
        neighborhood_rows,
        [2400, 700, 1900, 2200, 1870],
    )
    add_figure(
        doc, "03_mahalle_medyan_fiyat.png", "Şekil 4.2. Mahalle bazında medyan ilan fiyatları"
    )
    kruskal = summary["neighborhood_test"]
    add_para(
        doc,
        f"En az 20 ilan içeren 18 mahallede toplam n={tr_num(kruskal['n'])} gözlem karşılaştırılmıştır. "
        f"Kruskal-Wallis testi H={tr_num(kruskal['h_statistic'], 2)}, p<0,001 ve ε²="
        f"{tr_num(kruskal['epsilon_squared'], 3)} vermiştir. Buna göre mahalle fiyat dağılımlarının "
        "tamamının aynı olduğu yokluk hipotezi reddedilir. Epsilon kare yaklaşık 0,30 düzeyinde olup "
        "mahalle grupları ile fiyat sıraları arasında kayda değer bir ayrışmaya işaret eder.",
    )
    add_para(
        doc,
        "Güzelyalı'nın medyan toplam fiyatı ve medyan metrekare fiyatı yüksek görünürken Atakent'in "
        "medyan toplam fiyatında hem gözlem sayısı hem de brüt alan bileşimi etkili olabilir. Bu nedenle "
        "mahalle tablosu tek başına konumsal prim olarak yorumlanmamalı; alan ve diğer nitelikleri aynı "
        "anda kontrol eden hedonik sonuçlarla birlikte okunmalıdır.",
    )

    add_heading(doc, "4.3. Oda tipi ve sayısal ilişkiler", 2)
    add_figure(doc, "04_oda_fiyat_profili.png", "Şekil 4.3. Oda tipine göre fiyat profili")
    add_para(
        doc,
        "Oda sayısı yükseldikçe tipik toplam fiyatın arttığı görülmektedir; ancak oda sayısı alanla "
        "yüksek ölçüde bağlantılıdır. Bu nedenle grafikteki basit fark, oda eklemenin tek başına fiyatı "
        "artırdığı biçiminde okunamaz. Permütasyon öneminde net/brüt alan ve oda sayısının birlikte üst "
        "sıralarda bulunması da bu fiziksel büyüklük bileşimini desteklemektedir.",
    )
    add_figure(
        doc,
        "05_korelasyon_isiharitasi.png",
        "Şekil 4.4. Sayısal değişkenler için Spearman korelasyonları",
    )
    correlations = pd.read_csv(TABLES / "spearman_korelasyon.csv", index_col=0)
    price_corr = correlations["Fiyat (TL)"].drop("Fiyat (TL)").sort_values(ascending=False)
    add_para(
        doc,
        f"Fiyatla en yüksek sıra korelasyonları brüt alan (ρ={tr_num(price_corr['Brüt m²'], 3)}), net "
        f"alan (ρ={tr_num(price_corr['Net m²'], 3)}) ve oda sayısı (ρ={tr_num(price_corr['Oda Sayısı Numeric'], 3)}) "
        "için elde edilmiştir. Korelasyon matrisi ayrıca brüt ve net alanın birbirine çok yakın hareket "
        "ettiğini gösterir. Bu çoklu bağlantı, hedonik katsayılar yorumlanırken ve özellik önemleri "
        "karşılaştırılırken dikkate alınmıştır.",
    )

    add_heading(doc, "4.4. İkili grup karşılaştırmaları", 2)
    binary_rows = []
    for _, row in binary.iterrows():
        binary_rows.append(
            (
                row["variable"],
                f"{row['group_a']} / {row['group_b']}",
                f"{tl(row['median_a'])} / {tl(row['median_b'])}",
                tr_num(row["rank_biserial"], 3),
                "<0,001" if row["p_value_holm"] < 0.001 else tr_num(row["p_value_holm"], 3),
                "Evet" if bool(row["significant_0_05"]) else "Hayır",
            )
        )
    add_table(
        doc,
        "Tablo 4.3. İkili grup karşılaştırmaları",
        ["Değişken", "Gruplar", "Medyanlar", "rᵣᵦ", "Holm p", "Anlamlı"],
        binary_rows,
        [1750, 1750, 2500, 1000, 1000, 1070],
    )
    add_para(
        doc,
        "Site içinde olma, eşyalı olma ve balkon bulunması için Holm düzeltmesinden sonra fiyat "
        "dağılımları farklıdır. Site içindeki ilanların medyanı daha yüksek, eşyalı ilanların medyanı "
        "daha düşük ve balkonlu ilanların medyanı daha yüksektir. Eşyalı sonucu ters yönlü bir kalite "
        "etkisi olarak yorumlanmamalıdır; eşyalı ilanların küçük dairelerde veya farklı mahallelerde "
        "yoğunlaşması olasıdır. Asansör, kredi uygunluğu ve takas karşılaştırmaları yüzde 5 düzeyinde "
        "anlamlı değildir.",
    )

    add_heading(doc, "4.5. Hedonik model bulguları", 2)
    chosen_terms = [
        "ln_brut_m2",
        "bina_yasi",
        "banyo_sayisi",
        "oda_sayisi",
        "bulundugu_kat",
        "kat_sayisi",
        "ilan_gunu",
        "mahalle_Güzelyalı Mh.",
        "mahalle_Atakent Mh.",
        "site_Hayır",
    ]
    label_terms = {
        "ln_brut_m2": "ln(Brüt m²)",
        "bina_yasi": "Bina yaşı",
        "banyo_sayisi": "Banyo sayısı",
        "oda_sayisi": "Oda sayısı",
        "bulundugu_kat": "Bulunduğu kat",
        "kat_sayisi": "Toplam kat sayısı",
        "ilan_gunu": "İlan günü",
        "mahalle_Güzelyalı Mh.": "Güzelyalı (ref. Alanlı)",
        "mahalle_Atakent Mh.": "Atakent (ref. Alanlı)",
        "site_Hayır": "Site dışında (ref. site içinde)",
    }
    hedonic_rows = []
    for term in chosen_terms:
        row = hedonic.loc[hedonic["term"].eq(term)].iloc[0]
        percent_value = row["percent_change_per_unit"]
        percent_text = "Esneklik" if pd.isna(percent_value) else f"%{tr_num(percent_value, 1)}"
        hedonic_rows.append(
            (
                label_terms[term],
                tr_num(row["coefficient"], 4),
                tr_num(row["std_error_hc3"], 4),
                "<0,001" if row["p_value"] < 0.001 else tr_num(row["p_value"], 3),
                percent_text,
            )
        )
    add_table(
        doc,
        "Tablo 4.4. Hedonik modelin seçili katsayıları",
        ["Terim", "β", "HC3 SH", "p", "Yaklaşık fark"],
        hedonic_rows,
        [3000, 1400, 1400, 1200, 2070],
    )
    h = summary["hedonic_model"]
    add_para(
        doc,
        f"Model n={tr_num(h['n'])} ilan ve {tr_num(h['parameter_count'])} parametreyle tahmin edilmiştir. "
        f"R²={tr_num(h['r_squared'], 3)}, düzeltilmiş R²={tr_num(h['adjusted_r_squared'], 3)} ve log "
        f"ölçekte RMSE={tr_num(h['rmse_log'], 3)}'tür. Brüt alanın log-log katsayısı "
        f"{tr_num(hedonic.loc[hedonic['term'].eq('ln_brut_m2'), 'coefficient'].item(), 3)} olup, diğer "
        "değişkenler sabitken brüt alandaki yüzde 1 artışın ilan fiyatında yaklaşık yüzde 0,62 artışla "
        "birlikte gözlendiğini gösterir.",
    )
    add_para(
        doc,
        "Bina yaşındaki bir yıllık artış yaklaşık yüzde 1,8 daha düşük fiyatla; ilave banyo yaklaşık "
        "yüzde 10,5 ve bir kat daha yukarıda bulunma yaklaşık yüzde 3,4 daha yüksek fiyatla ilişkilidir. "
        "Oda sayısı, toplam kat sayısı ve ilan günü seçili modelde yüzde 5 düzeyinde anlamlı değildir. "
        "Güzelyalı ve Atakent göstergeleri Alanlı referansına göre pozitifken site dışında olma yaklaşık "
        "yüzde 6,3 daha düşük fiyatla ilişkilidir. Koşul sayısının yüksek olması, özellikle alan ve oda "
        "değişkenleri arasındaki bağlantı nedeniyle katsayıların temkinli yorumlanmasını gerektirir.",
    )

    add_heading(doc, "4.6. Model karşılaştırması", 2, new_page=True)
    comparison_rows = []
    for _, row in comparison.iterrows():
        comparison_rows.append(
            (
                row["model"],
                tl(row["cv_mae_mean"]),
                tl(row["cv_mae_std"]),
                tl(row["test_mae"]),
                tr_num(row["test_r2"], 3),
                pct(row["test_mape"]),
            )
        )
    add_table(
        doc,
        "Tablo 4.5. Tahmin modeli karşılaştırması",
        ["Model", "CV MAE", "CV SS", "Test MAE", "Test R²", "MAPE"],
        comparison_rows,
        [2500, 1600, 1400, 1500, 1000, 1070],
    )
    add_figure(
        doc, "06_model_performansi.png", "Şekil 4.5. Aday modellerin çapraz doğrulama performansı"
    )
    add_para(
        doc,
        f"En düşük eğitim içi CV MAE {model['name']} modelinde {tl(model['cross_validation']['cv_mae_mean'])} "
        f"olarak elde edilmiştir. Kilitli 385 ilanlık testte MAE {tl(model['locked_test']['test_mae'])}, "
        f"RMSE {tl(model['locked_test']['test_rmse'])}, R²={tr_num(model['locked_test']['test_r2'], 3)} ve "
        f"MAPE={pct(model['locked_test']['test_mape'])}'dir. Bootstrap yüzde 95 MAE aralığı "
        f"{tl(model['mae_ci_95_low'])}-{tl(model['mae_ci_95_high'])} olarak hesaplanmıştır.",
    )
    add_para(
        doc,
        "Histogram Gradient Boosting kilitli test R²'sinde çok küçük bir üstünlük gösterse de model "
        "seçim kuralı önceden belirlenen eğitim içi CV MAE'dir; bu nedenle test kümesine göre model "
        "değiştirilmemiştir. Medyan temel modelin negatif test R²'si, özellikleri kullanan tüm adayların "
        "anlamlı tahmin kazanımı sağladığını gösterir.",
    )

    add_heading(doc, "4.7. Tahmin hatalarının görünümü", 2, new_page=True)
    add_figure(
        doc,
        "07_gercek_tahmin.png",
        "Şekil 4.6. Kilitli testte gerçekleşen ve tahmin edilen fiyatlar",
    )
    add_para(
        doc,
        "Gerçek-tahmin grafiğinde noktaların 45 derece çizgisi çevresinde yoğunlaşması modelin genel "
        "fiyat sıralamasını yakaladığını gösterir. Bununla birlikte yüksek fiyatlı ilanlarda saçılım "
        "genişlemekte ve bazı değerler merkeze doğru çekilmektedir. Bu desen, ağaç topluluklarında "
        "örneklem sınırlarına yakın pahalı konutların daha zor tahmin edildiğini düşündürür.",
    )
    add_figure(doc, "08_artik_analizi.png", "Şekil 4.7. Kilitli test artıklarının dağılımı")
    add_para(
        doc,
        "Artık grafiği sıfır çevresinde belirgin bir yoğunluk gösterse de hata varyansı fiyat düzeyine "
        "göre sabit değildir. Bu nedenle tek bir MAE değeri her konut segmentinde eşit güven düzeyi "
        "anlamına gelmez. Dashboard simülatöründe raporlanan ±MAE ifadesi de bireysel güven aralığı "
        "olarak değil, testteki tipik mutlak hata referansı olarak etiketlenmiştir.",
    )

    add_heading(doc, "4.8. Özellik önemi", 2)
    add_figure(doc, "09_ozellik_onemi.png", "Şekil 4.8. Permütasyon özelliği önemleri")
    top_features = ", ".join(importance.head(5)["feature"].tolist())
    add_para(
        doc,
        f"Bir değişken testte karıştırıldığında MAE'nin ne kadar arttığını ölçen permütasyon analizinde "
        f"ilk beş özellik {top_features} olarak sıralanmıştır. Net alan karıştırıldığında MAE ortalama "
        f"{tl(importance.iloc[0]['mae_increase_mean'])}, brüt alan karıştırıldığında "
        f"{tl(importance.iloc[1]['mae_increase_mean'])} artmıştır. Korelasyonlu özelliklerin önemleri "
        "birbirinin yerine geçebileceğinden bu değerler nedensel katkı veya bağımsız fiyat primi değildir.",
    )

    add_heading(doc, "4.9. Değerlendirme protokolü duyarlılığı", 2)
    protocol_rows = []
    for _, row in protocols.iterrows():
        protocol_rows.append(
            (
                row["protocol"],
                tr_num(row["train_rows"]),
                tr_num(row["test_rows"]),
                tl(row["mae"]),
                tr_num(row["r2"], 3),
                pct(row["mape"]),
            )
        )
    add_table(
        doc,
        "Tablo 4.6. Değerlendirme protokolü duyarlılığı",
        ["Protokol", "Eğitim", "Test", "MAE", "R²", "MAPE"],
        protocol_rows,
        [3300, 900, 800, 1700, 1000, 1370],
    )
    add_figure(doc, "10_protokol_duyarliligi.png", "Şekil 4.9. Değerlendirme protokolü duyarlılığı")
    add_para(
        doc,
        "Tüm tekrar görüntülerin satır bazlı rastgele bölünmesi en iyi görünen sonucu üretmiştir "
        "(MAE 396.731 TL; R²=0,829). Aynı ilan numarasının eğitim ve testte birlikte bulunmasını "
        "engelleyen grup bölmesinde hata 524.420 TL'ye yükselmiş, R²=0,693'e düşmüştür. Bu fark, tekrar "
        "ilanların rastgele satır bölmesinde iyimserlik yarattığını doğrudan gösterir.",
    )
    add_para(
        doc,
        "Kronolojik kontrol 9 Mayıs 2025 kesimiyle MAE 502.892 TL ve R²=0,764 vermiştir. Ana kilitli "
        "tekil testin sonucu bu iki daha zor senaryonun arasında yer alır. Dolayısıyla modelin performansı "
        "tek bir yüzdeyle değil, kullanım senaryosuna bağlı yaklaşık R²=0,69-0,79 bandında değerlendirilmelidir.",
    )


def add_discussion(doc: Document, summary: dict) -> None:
    add_heading(doc, "5. TARTIŞMA", 1, new_page=True)
    add_heading(doc, "5.1. Temel bulguların birlikte değerlendirilmesi", 2)
    add_para(
        doc,
        "Bulgular üç ortak eksende birleşmektedir. Birincisi, konut büyüklüğü toplam ilan fiyatının en "
        "güçlü gözlenen belirleyici kümesidir. Brüt/net alan hem korelasyonlarda hem hedonik modelde hem "
        "de permütasyon öneminde öne çıkmıştır. İkincisi, mahalleler arasında basit dağılım karşılaştırması "
        "güçlü farklılık gösterirken, alan ve yapı özellikleri kontrol edildiğinde mahalle katsayılarının "
        "büyüklükleri değişmektedir. Üçüncüsü, model performansı veri bölme kuralına duyarlıdır.",
    )
    add_para(
        doc,
        "Bu üç sonuç, konut fiyatı analizinde tek değişkenli sıralamalarla yetinilmemesi gerektiğini "
        "gösterir. Örneğin toplam fiyatı yüksek bir mahallenin konutları daha büyük olabilir; balkonlu "
        "ilanlar daha yeni binalarda veya belirli konumlarda yoğunlaşabilir. Basit grup farkları piyasa "
        "profili sunarken, çok değişkenli model koşullu birliktelikleri ve makine öğrenmesi tahmin "
        "yapısını tamamlayıcı biçimde ele alır.",
    )

    add_heading(doc, "5.2. Literatürle ilişki", 2)
    add_para(
        doc,
        "Büyüklük ve konumun önemi, hedonik fiyat yaklaşımının temel beklentisiyle uyumludur (Rosen, "
        "1974). Doğrusal olmayan ağaç topluluklarının Ridge modelinden daha düşük hata üretmesi de konut "
        "nitelikleri ile fiyat arasındaki ilişkinin tek bir doğrusal yüzeyle tam temsil edilemeyebileceğini "
        "destekler. Türkiye verisinde hedonik regresyonla daha esnek tahmin yöntemlerini karşılaştıran "
        "Selim'in (2009) yaklaşımıyla yöntemsel bir paralellik vardır; ancak veri kaynağı ve hedef değişken "
        "farklı olduğu için sayıların doğrudan kıyaslanması doğru değildir.",
    )
    add_para(
        doc,
        "Extra Trees'in seçilmesi, rastgele eşiklere dayalı ağaç topluluğunun bu veri kümesindeki karmaşık "
        "etkileşimleri yakalayabildiğini gösterir. Bununla birlikte Histogram Gradient Boosting'in test "
        "sonuçları çok yakındır. Bu yakınlık, algoritma adından çok veri kalitesinin, protokolün ve "
        "belirsizlik raporlamasının önemini ortaya koymaktadır.",
    )

    add_heading(doc, "5.3. Önceki sürümdeki performans iddialarının yeniden okunması", 2)
    add_para(
        doc,
        "Önceki sürümde farklı veri kesitleri ve modellemeler için daha yüksek R² değerleri raporlanmıştır. "
        "Yenilenmiş analiz bu değerleri geçersiz ilan etmek yerine, doğrudan yeniden üretilemeyen geçmiş "
        "sonuçlar olarak ayırır. Mevcut veriyle satır bazlı rastgele bölmenin R²'yi 0,829'a çıkarması, aynı "
        "ilanın tekrar görüntülerinin performansı yükseltebildiğini göstermektedir. Bu nedenle yeni ana "
        "sonuç; tekil ilan, eğitim içinde model seçimi ve kilitli test ilkelerine dayanır.",
    )
    add_para(
        doc,
        "Model başarısının daha temkinli raporlanması bir performans kaybı değil, ölçüm tanımının "
        "iyileştirilmesidir. Portföy veya akademik kullanımda R²=0,79 değerinin hangi test büyüklüğü, "
        "hangi veri birimi ve hangi rastgelelik tohumu altında elde edildiğinin gösterilmesi, kaynağı "
        "belirsiz daha yüksek bir orandan daha değerlidir.",
    )

    add_heading(doc, "5.4. Sınırlılıklar", 2)
    limitations = [
        (
            "Örnekleme",
            "İlanlar olasılıklı örneklem değildir; platform ve toplama tercihlerine bağlıdır.",
        ),
        ("Hedef", "Fiyat, gerçekleşmiş tapu satış bedeli değil satıcının ilan talebidir."),
        (
            "Ham veri",
            "İlk 2.836 satırlık dosya olmadığı için önceden çıkarılan 555 kaydın gerekçesi doğrulanamaz.",
        ),
        (
            "Zaman",
            "Gözlemler büyük ölçüde Nisan-Mayıs 2025'te yoğunlaşır; güçlü zaman serisi çıkarımı uygun değildir.",
        ),
        (
            "Eksik nitelik",
            "Manzara, iç kalite, bina konumu ve mikro çevre gibi fiyatla ilişkili özellikler sınırlıdır.",
        ),
        (
            "Mekânsal yapı",
            "Koordinat bulunmadığından mekânsal otokorelasyon ve mesafe etkileri modellenmemiştir.",
        ),
        (
            "Belirsizlik",
            "MAE aralığı bireysel tahmin aralığı değildir; simülatör resmi ekspertiz sunmaz.",
        ),
    ]
    add_table(
        doc,
        "Sınırlılıkların sonuçlara etkisi",
        ["Alan", "Sınır"],
        limitations,
        [1800, 7270],
        source="Kaynak: Araştırmanın kapsam değerlendirmesi.",
    )

    add_heading(doc, "5.5. Uygulama değeri", 2)
    add_para(
        doc,
        "Dinamik dashboard, sabit tez tablolarını tamamlayan bir keşif katmanıdır. Kullanıcı mahalle, oda "
        "tipi, site durumu, fiyat ve alan aralığıyla piyasayı filtreleyebilir; model karşılaştırması ve "
        "protokol duyarlılığını inceleyebilir. Fiyat simülatörü seçilen senaryo için model tahmini ile aynı "
        "mahalle/oda/alan bandındaki ilanların medyanını yan yana verir. Veri kalite sekmesi 2.836'dan "
        "1.922 tekil ilana uzanan akışı görünür kılar.",
    )
    add_para(
        doc,
        "Bu yapı bir emlak değerleme ürünü olmaktan çok, istatistiksel düşünme ve model yönetişimi "
        "gösterimidir. Bir sonraki aşamada gerçek satış verisi, koordinat, ulaşım/deniz mesafesi ve düzenli "
        "zaman güncellemeleri eklendiğinde hem hedonik yorum hem de tahmin geçerliliği güçlenebilir.",
    )


def add_conclusion(doc: Document, summary: dict) -> None:
    model = summary["best_model"]
    add_heading(doc, "6. SONUÇ VE ÖNERİLER", 1, new_page=True)
    add_heading(doc, "6.1. Sonuç", 2)
    add_para(
        doc,
        "Atakum satılık konut ilanlarının yenilenmiş analizi, ilan fiyatlarının sağa çarpık olduğunu; "
        "alan, oda ve banyo özellikleri ile mahalle bileşiminin fiyat farklılıklarında önemli rol "
        "oynadığını göstermiştir. Uygun mahalleler arasında fiyat dağılımları güçlü biçimde ayrışmış, "
        "yarı logaritmik hedonik model gözlenen log-fiyat varyansının yaklaşık yüzde 79'unu açıklamıştır. "
        "Bina yaşı negatif; brüt alan, banyo sayısı ve bulunduğu kat pozitif koşullu ilişkiler vermiştir.",
    )
    add_para(
        doc,
        f"Tahmin tarafında {model['name']} eğitim içi çapraz doğrulamayla seçilmiş ve kilitli testte "
        f"{tl(model['locked_test']['test_mae'])} MAE ile R²={tr_num(model['locked_test']['test_r2'], 3)} "
        "elde etmiştir. Daha önemli sonuç, aynı ilan görüntülerini satır bazında rastgele ayıran protokolün "
        "hata değerini yapay biçimde düşürebilmesidir. Model başarısı ancak gözlem birimi ve kullanım "
        "senaryosuyla birlikte anlamlıdır.",
    )

    add_heading(doc, "6.2. Öneriler", 2)
    recommendations = [
        (
            "Veri",
            "Ham ve temiz veri sürümlerini ayrı tutun; her satır silme kararını neden koduyla kaydedin.",
        ),
        (
            "Kimlik",
            "İlan numarasını zorunlu grup anahtarı yapın; tekrar görüntüleri zaman boyutunda saklayın.",
        ),
        ("Hedef", "Mümkünse tapu/satış verisiyle ilan-satış farkını kalibre edin."),
        (
            "Konum",
            "Koordinat, kıyı/ulaşım/üniversite mesafesi ve mekânsal çapraz doğrulama ekleyin.",
        ),
        ("Zaman", "Düzenli veri toplama ile fiyat değişimi ve ilan süresi modelleri geliştirin."),
        (
            "Model",
            "Tek sonuç yerine kilitli, kronolojik ve grup bölmeli performansları birlikte raporlayın.",
        ),
        (
            "Ürün",
            "Dashboard tahminlerini veri dönemi, örnek sayısı ve kullanım sınırıyla birlikte gösterin.",
        ),
    ]
    add_table(
        doc,
        "Gelecek çalışma önerileri",
        ["Başlık", "Öneri"],
        recommendations,
        [1700, 7370],
        source="Kaynak: Bulgulara dayalı araştırma önerileri.",
    )

    add_heading(doc, "6.3. Nihai değerlendirme", 2)
    add_para(
        doc,
        "Çalışmanın temel çıktısı yalnızca bir fiyat tahmin modeli değildir. Veri soy ağacı, sağlam "
        "istatistiksel testler, hedonik açıklama, kilitli tahmin değerlendirmesi, dinamik dashboard ve "
        "sürüm kontrollü kod birlikte bir araştırma ürünü oluşturur. Bu bütünlük, sonucun yeniden "
        "üretilebilmesini ve farklı varsayımların etkisinin görülebilmesini sağlar.",
    )


def add_references(doc: Document) -> None:
    add_heading(doc, "KAYNAKÇA", 1, new_page=True)
    references = [
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32. https://doi.org/10.1023/A:1010933404324",
        "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189-1232. https://doi.org/10.1214/aos/1013203451",
        "Geurts, P., Ernst, D., & Wehenkel, L. (2006). Extremely randomized trees. Machine Learning, 63, 3-42. https://doi.org/10.1007/s10994-006-6226-1",
        "Hoerl, A. E., & Kennard, R. W. (1970). Ridge regression: Biased estimation for nonorthogonal problems. Technometrics, 12(1), 55-67. https://doi.org/10.1080/00401706.1970.10488634",
        "Holm, S. (1979). A simple sequentially rejective multiple test procedure. Scandinavian Journal of Statistics, 6(2), 65-70.",
        "Kruskal, W. H., & Wallis, W. A. (1952). Use of ranks in one-criterion variance analysis. Journal of the American Statistical Association, 47(260), 583-621. https://doi.org/10.1080/01621459.1952.10483441",
        "MacKinnon, J. G., & White, H. (1985). Some heteroskedasticity-consistent covariance matrix estimators with improved finite sample properties. Journal of Econometrics, 29(3), 305-325. https://doi.org/10.1016/0304-4076(85)90158-7",
        "Mann, H. B., & Whitney, D. R. (1947). On a test of whether one of two random variables is stochastically larger than the other. The Annals of Mathematical Statistics, 18(1), 50-60. https://doi.org/10.1214/aoms/1177730491",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Rosen, S. (1974). Hedonic prices and implicit markets: Product differentiation in pure competition. Journal of Political Economy, 82(1), 34-55. https://doi.org/10.1086/260169",
        "Selim, H. (2009). Determinants of house prices in Turkey: Hedonic regression versus artificial neural network. Expert Systems with Applications, 36(2), 2843-2852. https://doi.org/10.1016/j.eswa.2008.01.044",
        "Sahibinden.com. (2025). Satılık konut ilanları [Araştırma kapsamında manuel olarak derlenen çevrim içi ilan verisi].",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Kaynakça")
        run = paragraph.add_run(reference)
        set_run_font(run, size=10)


def add_appendices(doc: Document) -> None:
    add_heading(doc, "EKLER", 1, new_page=True)
    add_heading(doc, "Ek A. Tekrarlanabilir çalıştırma", 2)
    add_para(
        doc,
        "Projenin kök dizininde Python 3.10 veya üzeriyle aşağıdaki komutlar çalıştırılır. İlk komut "
        "bağımlılıkları, ikinci komut analiz tablolarını ve şekilleri, üçüncü komut dashboard'u başlatır.",
    )
    commands = [
        "python -m pip install -r requirements.txt",
        "python analiz_icin_kodlar.py",
        "streamlit run dashboard/app.py",
        "python -m pytest",
    ]
    for command in commands:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.right_indent = Cm(0.4)
        paragraph.paragraph_format.space_after = Pt(4)
        set_cell_like = OxmlElement("w:shd")
        set_cell_like.set(qn("w:fill"), LIGHT_BLUE)
        paragraph._p.get_or_add_pPr().append(set_cell_like)
        run = paragraph.add_run(command)
        set_run_font(run, name="Consolas", size=9.5, color=NAVY)

    add_heading(doc, "Ek B. Başlıca çıktı dosyaları", 2)
    output_rows = [
        ("outputs/latest/ANALIZ_OZETI.md", "İnsan tarafından okunabilir ana sonuçlar"),
        ("outputs/latest/analiz_ozeti.json", "Dashboard ve rapor için makine okunur özet"),
        ("outputs/latest/tables/", "İstatistik, model ve sağlamlık tabloları"),
        ("outputs/latest/figures/", "Tez ve README görselleri"),
        ("dashboard/app.py", "Dört sekmeli dinamik dashboard ve fiyat simülatörü"),
        ("src/atakum_housing/", "Veri, istatistik, modelleme ve görselleştirme paketi"),
        ("tests/", "Veri hattı, metrik ve dashboard testleri"),
    ]
    add_table(
        doc,
        "Repo çıktı haritası",
        ["Yol", "İçerik"],
        output_rows,
        [3700, 5370],
        source="Kaynak: Proje deposu.",
    )

    add_heading(doc, "Ek C. Dashboard kullanım sınırı", 2)
    add_para(
        doc,
        "Dashboard'daki filtreler tekil son ilan örneklemini değiştirir. Küçük filtre grupları piyasa "
        "genelini temsil etmeyebilir. Simülatör, bütün tekil ilanlar üzerinde eğitilen Histogram Gradient "
        "Boosting modeliyle hızlı senaryo üretir; tezde model seçimi için raporlanan Extra Trees kilitli "
        "test sonucunun yerine geçmez. Benzer ilan medyanı yalnızca mahalle, oda ve ±%20 alan aralığına "
        "dayanır. Her sonuç, gerçekleşmiş satış fiyatı değil ilan fiyatı bağlamında değerlendirilmelidir.",
    )


def build(output: Path) -> Path:
    summary = load_json("analiz_ozeti.json")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc, summary)
    add_contents(doc)
    add_introduction(doc)
    add_literature(doc)
    add_methodology(doc, summary)
    add_results(doc, summary)
    add_discussion(doc, summary)
    add_conclusion(doc, summary)
    add_references(doc)
    add_appendices(doc)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
    doc.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    print(build(args.output))


if __name__ == "__main__":
    main()
